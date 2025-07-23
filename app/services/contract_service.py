#Genera un contrato Word a partir del template y los datos recibidos.

import os
from docx import Document
from jinja2 import Environment
from datetime import datetime

# Ruta a las carpetas de templates y de salida (output)
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), '..', 'templates')
OUTPUT_DIR = os.path.join(TEMPLATES_DIR, 'output')

# Asegura que la carpeta de salida exista
os.makedirs(OUTPUT_DIR, exist_ok=True)

def render_contract(context: dict, template_filename: str = "prestacion_servicios.docx") -> str:
    """
    Genera un contrato Word rellenando un template con los datos entregados.
    :param context: Diccionario con los campos y valores del contrato.
    :param template_filename: Nombre del template base en /templates/.
    :return: Ruta absoluta al archivo generado listo para descargar.
    """
    # Ruta del template base
    template_path = os.path.join(TEMPLATES_DIR, template_filename)
    # Nombre único para el archivo generado
    output_filename = f"Contrato_{context['nombre_contratante']}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)

    # Carga el documento Word base
    doc = Document(template_path)

    # Usamos Jinja2 para reemplazo de variables en los textos tipo {{campo}}
    env = Environment()

    # Procesa todos los párrafos (texto libre)
    for p in doc.paragraphs:
        inline = p.runs
        for i in range(len(inline)):
            text = inline[i].text
            if "{{" in text and "}}" in text:
                # Si encuentra una variable, la reemplaza con Jinja2
                template = env.from_string(text)
                inline[i].text = template.render(context)

    # Procesa tablas, si existen (pueden tener variables también)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                if "{{" in text and "}}" in text:
                    template = env.from_string(text)
                    cell.text = template.render(context)

    # Guarda el archivo Word generado en la carpeta /output
    doc.save(output_path)
    return output_path
